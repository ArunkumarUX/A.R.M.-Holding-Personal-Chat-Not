#!/usr/bin/env python3
"""Generate Word doc from ADGM-Agent-Behaviours-and-Answer-Formats.md"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "docs" / "ADGM-Agent-Behaviours-and-Answer-Formats.md"
OUT = ROOT / "docs" / "ADGM-Agent-Behaviours-and-Answer-Formats.docx"


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx < len(headers):
                table.rows[r_idx + 1].cells[c_idx].text = val


def parse_table_row(line: str) -> list[str]:
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def is_separator_row(line: str) -> bool:
    s = line.strip().strip("|").replace(" ", "")
    return s and all(c in "-:" for c in s)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def build_from_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    i = 0
    table_headers: list[str] | None = None
    table_rows: list[list[str]] = []
    code_lines: list[str] | None = None
    bullet_buf: list[str] = []

    def flush_bullets() -> None:
        nonlocal bullet_buf
        for item in bullet_buf:
            doc.add_paragraph(item, style="List Bullet")
        bullet_buf = []

    while i < len(lines):
        line = lines[i]

        if code_lines is not None:
            if line.strip().startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = None
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_bullets()
            code_lines = []
            i += 1
            continue

        if line.strip().startswith("|") and "|" in line[1:]:
            flush_bullets()
            if is_separator_row(line):
                i += 1
                continue
            cells = parse_table_row(line)
            if table_headers is None:
                table_headers = cells
                table_rows = []
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if table_headers is not None:
                add_table(doc, table_headers, table_rows)
                table_headers = None
                table_rows = []

        if line.startswith("# ") and not line.startswith("## "):
            flush_bullets()
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith("## "):
            flush_bullets()
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            flush_bullets()
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("#### "):
            flush_bullets()
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        if line.strip() == "---":
            flush_bullets()
            doc.add_paragraph()
            i += 1
            continue

        if line.strip().startswith("- ") or line.strip().startswith("* "):
            bullet_buf.append(line.strip()[2:].strip())
            i += 1
            continue

        if line.strip().startswith("**") and line.strip().endswith("**") and ":" in line:
            flush_bullets()
            p = doc.add_paragraph()
            raw = line.strip()
            parts = raw.split(":", 1)
            label = parts[0].strip("* ").strip()
            rest = parts[1].strip() if len(parts) > 1 else ""
            r0 = p.add_run(label + ": ")
            r0.bold = True
            if rest:
                p.add_run(rest)
            i += 1
            continue

        if line.strip():
            flush_bullets()
            p = doc.add_paragraph(line.strip())
            if line.strip().startswith("*") and line.strip().endswith("*"):
                for run in p.runs:
                    run.italic = True
        i += 1

    flush_bullets()
    if table_headers is not None:
        add_table(doc, table_headers, table_rows)
    if code_lines:
        add_code_block(doc, code_lines)


def main() -> None:
    if not MD.exists():
        raise SystemExit(f"Missing source: {MD}")

    doc = Document()
    set_doc_defaults(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    body = MD.read_text(encoding="utf-8")
    build_from_markdown(doc, body)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
